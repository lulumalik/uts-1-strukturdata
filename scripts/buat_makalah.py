# -*- coding: utf-8 -*-
"""Generate makalah DOCX: cover dari file lama, isi disesuaikan app terkini."""

from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "TUGAS SDA TERBARU.docx")
OUT = os.path.join(ROOT, "MAKALAH_SDA_EHEALTH_TERBARU.docx")


def set_run_font(run, name="Times New Roman", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_custom(doc, text, size=14, center=True, space_before=18, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_subheading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


_bullet_n = 0
_number_n = 0


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("• " + text)
    set_run_font(run, size=12)
    return p


def add_numbered(doc, text):
    global _number_n
    _number_n += 1
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{_number_n}. {text}")
    set_run_font(run, size=12)
    return p


def reset_numbering():
    global _number_n
    _number_n = 0


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=11)
    doc.add_paragraph()
    return table


def keep_cover_only(doc):
    """Hapus semua konten setelah cover (sebelum KATA PENGANTAR)."""
    body = doc.element.body
    children = list(body)
    cut_at = None
    for i, child in enumerate(children):
        texts = [t.text or "" for t in child.iter(qn("w:t"))]
        joined = "".join(texts).strip().upper()
        if joined == "KATA PENGANTAR":
            cut_at = i
            break
    if cut_at is None:
        raise RuntimeError("KATA PENGANTAR tidak ditemukan di cover source")

    # Hapus dari KATA PENGANTAR sampai sebelum sectPr (elemen terakhir biasanya sectPr)
    for child in children[cut_at:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def build_content(doc):
    # ========== KATA PENGANTAR ==========
    add_heading_custom(doc, "KATA PENGANTAR", size=14)

    add_body(
        doc,
        'Puji syukur ke hadirat Tuhan Yang Maha Esa atas segala rahmat, karunia, dan '
        'penyertaan-Nya sehingga penulis dapat menyelesaikan makalah yang berjudul '
        '"Implementasi Struktur Data dan Algoritma pada Sistem Manajemen Antrian dan '
        'Rekam Medis Pasien Berbasis E-Health Menggunakan Bahasa Pemrograman C++" dengan baik.',
    )
    add_body(
        doc,
        "Makalah ini disusun sebagai salah satu syarat untuk memenuhi tugas pada mata kuliah "
        "Struktur Data dan Algoritma. Penyusunan makalah bertujuan mengimplementasikan konsep "
        "struktur data dan algoritma ke dalam aplikasi Sistem Manajemen Antrian dan Rekam Medis "
        "Pasien Berbasis E-Health yang bermanfaat bagi petugas klinik maupun pembelajaran mahasiswa.",
    )
    add_body(
        doc,
        "Dalam implementasinya, sistem memanfaatkan tipe data dan variabel (struct Pasien), "
        "Array, Linked List, Queue, Stack, Tree, dan Graph, serta algoritma Linear Search, "
        "Binary Search, Bubble Sort, Selection Sort, Insertion Sort, Breadth First Search (BFS), "
        "dan Depth First Search (DFS). Data pasien juga disimpan secara permanen menggunakan "
        "SQLite (file lokal data/klinik.db) agar tidak hilang saat program ditutup.",
    )
    add_body(
        doc,
        "Penulis menyadari bahwa makalah ini masih memiliki keterbatasan. Kritik dan saran yang "
        "bersifat membangun sangat diharapkan demi penyempurnaan makalah dan pengembangan "
        "aplikasi di masa yang akan datang.",
    )
    add_body(
        doc,
        "Akhir kata, penulis berharap makalah ini dapat memberikan manfaat serta menambah "
        "wawasan mengenai penerapan struktur data dan algoritma dalam sistem informasi "
        "pelayanan kesehatan.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Malang, Juli 2026\n\nPenyusun")
    set_run_font(run, size=12)

    for nama in [
        "1. LULU MAULANA MALIK (250401010517)",
        "2. A. M. NABIL AIENDRA DZIKRA (250401010507)",
        "3. MUHAMMAD FARIS (250401010509)",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(nama)
        set_run_font(run, size=12)

    # ========== BAB I ==========
    add_heading_custom(doc, "BAB I", size=14)
    add_heading_custom(doc, "PENDAHULUAN", size=14, space_before=0)

    add_subheading(doc, "1.1 Latar Belakang")
    add_body(
        doc,
        "Perkembangan teknologi informasi di bidang kesehatan mendorong digitalisasi pelayanan "
        "klinik agar lebih efisien, akurat, dan cepat. Pengelolaan data pasien dan antrian yang "
        "masih manual sering menimbulkan duplikasi nomor rekam medis, kesulitan pencarian data, "
        "serta antrian yang tidak teratur.",
    )
    add_body(
        doc,
        "Dengan kapasitas sekitar 150 pasien per bulan, diperlukan sistem yang mampu mengelola "
        "data secara terstruktur. Penelitian ini mengembangkan Sistem Manajemen Antrian dan "
        "Rekam Medis Pasien Berbasis E-Health menggunakan C++. Sistem mengimplementasikan "
        "Array, Linked List, Queue, Stack, Tree, dan Graph, serta algoritma searching, sorting, "
        "BFS, dan DFS. Persistensi data pasien menggunakan SQLite agar data tetap tersedia "
        "setelah aplikasi ditutup.",
    )

    add_subheading(doc, "1.2 Rumusan Masalah")
    add_body(doc, "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah:", first_line_indent=False)
    reset_numbering()
    add_numbered(doc, "Bagaimana implementasi tipe data, variabel, Array, Linked List, Queue, Stack, dan Tree pada sistem e-health?")
    add_numbered(doc, "Bagaimana penerapan Graph beserta BFS dan DFS untuk jaringan rujukan antar poli?")
    add_numbered(doc, "Bagaimana penerapan sorting (Bubble, Selection, Insertion) serta searching (Linear dan Binary) disertai analisis Big O?")
    add_numbered(doc, "Bagaimana sistem mengelola data pasien, antrian, riwayat medis, dan persistensi data secara efisien?")

    add_subheading(doc, "1.3 Tujuan Penelitian")
    reset_numbering()
    add_numbered(doc, "Mengimplementasikan struktur data dan algoritma pada Sistem Manajemen Antrian dan Rekam Medis Pasien berbasis e-health.")
    add_numbered(doc, "Membangun fitur CRUD pasien, antrian (Queue), riwayat (Stack), validasi poli (Tree), graph rujukan (BFS/DFS), sorting, dan searching dengan analisis Big O.")
    add_numbered(doc, "Menyediakan penyimpanan data pasien secara permanen menggunakan SQLite.")
    add_numbered(doc, "Menganalisis manfaat penerapan konsep tersebut bagi pengguna di bidang pelayanan kesehatan.")

    add_subheading(doc, "1.4 Batasan Masalah")
    add_bullet(doc, "Sistem dibangun dengan bahasa C++ (C++17) sebagai aplikasi console (CLI).")
    add_bullet(doc, "Kapasitas data pasien maksimal 150 per bulan (MAX_PASIEN = 150).")
    add_bullet(doc, "Poli yang tersedia: Umum, Gigi, Anak, dan Penyakit Dalam.")
    add_bullet(doc, "Data pasien disimpan di SQLite (data/klinik.db); antrian dan stack riwayat masih bersifat sesi program (in-memory).")
    add_bullet(doc, "Sistem merupakan simulasi administrasi klinik, bukan proses pemeriksaan medis nyata.")
    add_bullet(doc, "Tidak terdapat autentikasi pengguna maupun antarmuka GUI.")

    add_subheading(doc, "1.5 Manfaat Penelitian")
    add_body(doc, "Manfaat teoritis: memperdalam pemahaman implementasi struktur data dan algoritma pada kasus nyata e-health.", first_line_indent=False)
    add_body(doc, "Manfaat praktis: membantu petugas mengelola rekam medis, antrian FIFO, riwayat LIFO, validasi poli, rujukan antar poli, serta pengurutan dan pencarian data pasien.", first_line_indent=False)
    add_body(doc, "Manfaat akademik: memenuhi tugas mata kuliah Struktur Data dan Algoritma serta menjadi bahan laporan dan presentasi kelompok.", first_line_indent=False)

    # ========== BAB II ==========
    add_heading_custom(doc, "BAB II", size=14)
    add_heading_custom(doc, "LANDASAN TEORI", size=14, space_before=0)

    add_subheading(doc, "2.1 Tipe Data dan Variabel")
    add_body(
        doc,
        "Tipe data dan variabel menentukan cara data direpresentasikan dalam program. Pada sistem "
        "ini digunakan struct Pasien sebagai tipe data terstruktur yang merepresentasikan rekam "
        "medis pasien, serta konstanta MAX_PASIEN = 150 sebagai batasan kapasitas.",
    )
    add_table(
        doc,
        ["Field", "Tipe", "Keterangan"],
        [
            ["noRM", "string", "Nomor rekam medis (unik / primary key)"],
            ["nama", "string", "Nama lengkap pasien"],
            ["umur", "int", "Usia pasien"],
            ["poli", "string", "Poli tujuan"],
            ["status", "string", "Menunggu / Diperiksa / Selesai"],
        ],
    )

    add_subheading(doc, "2.2 Array")
    add_body(
        doc,
        "Array adalah struktur data statis berindeks. Akses elemen bersifat O(1). Dalam aplikasi, "
        "kelas ArrayPasien menyimpan data pasien hingga 150 elemen dan menjadi target utama "
        "algoritma sorting serta binary search.",
    )

    add_subheading(doc, "2.3 Linked List Linear")
    add_body(
        doc,
        "Linked list linear terdiri atas node yang saling terhubung melalui pointer. Kelas "
        "LinkedListPasien digunakan untuk penyimpanan dinamis: insert di head, pencarian "
        "sequential, update, dan hapus dengan manipulasi pointer. Kompleksitas pencarian O(n).",
    )

    add_subheading(doc, "2.4 Queue (FIFO)")
    add_body(
        doc,
        "Queue menerapkan prinsip First In First Out. Kelas QueueAntrian merepresentasikan "
        "antrian pelayanan klinik: enqueue saat pasien masuk antrian, dequeue saat dilayani. "
        "Manfaatnya adalah pelayanan yang adil sesuai urutan kedatangan.",
    )

    add_subheading(doc, "2.5 Stack (LIFO)")
    add_body(
        doc,
        "Stack menerapkan prinsip Last In First Out. Kelas StackRiwayat menyimpan riwayat "
        "pemeriksaan: push otomatis setelah dequeue, pop untuk mengambil riwayat terbaru. "
        "Hal ini memudahkan petugas melihat catatan pemeriksaan terakhir.",
    )

    add_subheading(doc, "2.6 Tree")
    add_body(
        doc,
        "Tree merepresentasikan hierarki. TreePoli memodelkan struktur Klinik E-Health dengan "
        "anak: Poli Umum, Poli Gigi, Poli Anak, dan Poli Penyakit Dalam. Pencarian poli "
        "dilakukan secara rekursif (child dan sibling) untuk memvalidasi input pasien.",
    )

    add_subheading(doc, "2.7 Graph, BFS, dan DFS")
    add_body(
        doc,
        "Graph terdiri atas vertex dan edge. Pada aplikasi, graph merepresentasikan hubungan "
        "rujukan antar poli. BFS menelusuri simpul berdasarkan level (lebar), sedangkan DFS "
        "menelusuri jalur sedalam mungkin. Kompleksitas keduanya O(V + E).",
    )

    add_subheading(doc, "2.8 Sorting, Searching, dan Big O")
    add_body(
        doc,
        "Sorting yang diimplementasikan meliputi Bubble Sort, Selection Sort, dan Insertion Sort "
        "dengan kriteria No RM, Nama, atau Umur (umumnya O(n²)). Searching mencakup Linear Search "
        "O(n) dan Binary Search O(log n) yang memerlukan data terurut berdasarkan No RM. "
        "Notasi Big O digunakan untuk menganalisis efisiensi operasi pada sistem.",
    )

    add_subheading(doc, "2.9 Persistensi SQLite")
    add_body(
        doc,
        "SQLite digunakan sebagai basis data lokal ringan. Setiap operasi tambah, ubah, dan hapus "
        "pasien disimpan ke file data/klinik.db. Saat program dibuka kembali, data dimuat otomatis "
        "ke Array dan Linked List sehingga pengguna tidak perlu menginput ulang.",
    )

    add_subheading(doc, "2.10 Ringkasan Kompleksitas")
    add_table(
        doc,
        ["Operasi", "Struktur / Algoritma", "Big O"],
        [
            ["Akses indeks", "Array", "O(1)"],
            ["Tambah di head", "Linked List", "O(1)"],
            ["Cari pasien", "Linear Search", "O(n)"],
            ["Cari No RM terurut", "Binary Search", "O(log n)"],
            ["Enqueue", "Queue", "O(1)"],
            ["Dequeue (array shift)", "Queue", "O(n)"],
            ["Push / Pop", "Stack", "O(1)"],
            ["Cari poli", "Tree (rekursif)", "O(n)"],
            ["Sorting", "Bubble / Selection / Insertion", "O(n²)"],
            ["BFS / DFS", "Graph", "O(V + E)"],
        ],
    )

    # ========== BAB III ==========
    add_heading_custom(doc, "BAB III", size=14)
    add_heading_custom(doc, "ANALISIS DAN PERANCANGAN SISTEM", size=14, space_before=0)

    add_subheading(doc, "3.1 Analisis Sistem")
    add_body(
        doc,
        "Pengelolaan manual pada klinik dengan kapasitas hingga 150 pasien per bulan menimbulkan "
        "risiko duplikasi data, kesalahan input, antrian tidak teratur, sulitnya melacak riwayat, "
        "serta belum adanya model rujukan antar poli. Sistem e-health ini dirancang untuk menjawab "
        "kebutuhan tersebut melalui struktur data dan algoritma yang tepat.",
    )

    add_subheading(doc, "3.2 Pemetaan Kebutuhan ke Struktur Data")
    add_table(
        doc,
        ["Kebutuhan Operasional", "Struktur Data / Algoritma"],
        [
            ["Simpan dan akses data pasien", "Array + Linked List + SQLite"],
            ["Antrian pelayanan", "Queue (FIFO)"],
            ["Riwayat pemeriksaan", "Stack (LIFO)"],
            ["Validasi poli", "Tree"],
            ["Jaringan rujukan poli", "Graph + BFS + DFS"],
            ["Daftar pasien terurut", "Bubble / Selection / Insertion Sort"],
            ["Cari No RM + analisis efisiensi", "Linear Search dan Binary Search + Big O"],
        ],
    )

    add_subheading(doc, "3.3 Kebutuhan Fungsional (Menu Aplikasi)")
    add_table(
        doc,
        ["Menu", "Fitur", "Konsep SDA"],
        [
            ["1", "Tambah Data Pasien", "Array, Linked List, Tree, SQLite"],
            ["2", "Cari Data Pasien", "Linked List (sequential search)"],
            ["3–4", "Update / Hapus Pasien", "Array, Linked List, SQLite"],
            ["5–6", "Tampil Linked List / Array", "Traversal linear"],
            ["7", "Kelola Antrian", "Queue enqueue / dequeue"],
            ["8", "Riwayat Medis", "Stack push / pop"],
            ["9", "Struktur Poli Klinik", "Tree"],
            ["10", "Statistik Pasien", "Agregasi data"],
            ["11", "Sorting Data Pasien", "Bubble, Selection, Insertion"],
            ["12", "Searching dan Big O", "Linear, Binary, analisis Big O"],
            ["13–15", "Graph, BFS, DFS", "Graph traversal"],
            ["0", "Keluar", "Tutup database"],
        ],
    )

    add_subheading(doc, "3.4 Perancangan Data Pasien")
    add_body(
        doc,
        "Setiap pasien memiliki No RM unik. Saat ditambahkan, data divalidasi (tidak kosong, "
        "tidak duplikat, poli valid menurut Tree), lalu disimpan ke Linked List, Array, dan "
        "SQLite secara bersamaan agar konsisten.",
    )

    add_subheading(doc, "3.5 Alur Sistem")
    add_body(
        doc,
        "Alur utama: Input data → Validasi (kosong / duplikat / poli) → Simpan ke Array + "
        "Linked List + SQLite → (opsional) Enqueue antrian → Dequeue pelayanan → Push riwayat "
        "Stack → Update status. Untuk laporan: Sorting pada Array; Searching Linear/Binary "
        "dengan tampilan Big O; Graph BFS/DFS untuk penelusuran rujukan poli.",
        first_line_indent=False,
    )

    add_subheading(doc, "3.6 Error Handling")
    add_bullet(doc, "Data kosong ditolak.")
    add_bullet(doc, "Duplikasi No RM ditolak.")
    add_bullet(doc, "Poli tidak valid ditolak berdasarkan Tree.")
    add_bullet(doc, "Array penuh memicu rollback Linked List.")
    add_bullet(doc, "Queue/Stack kosong atau penuh ditangani dengan pesan error.")
    add_bullet(doc, "Kegagalan simpan SQLite membatalkan penambahan data di memori.")

    add_subheading(doc, "3.7 Struktur File Project")
    add_body(
        doc,
        "Project modular: folder include/ untuk header (types, array, linkedlist, queue, stack, "
        "tree, graph, sortsearch, database), folder src/ untuk implementasi, third_party/sqlite/ "
        "untuk mesin SQLite, serta build.bat / Makefile untuk kompilasi. Output program: "
        "bin/e-health-klinik.exe.",
        first_line_indent=False,
    )

    # ========== BAB IV ==========
    add_heading_custom(doc, "BAB IV", size=14)
    add_heading_custom(doc, "IMPLEMENTASI DAN PEMBAHASAN", size=14, space_before=0)

    add_subheading(doc, "4.1 Lingkungan Implementasi")
    add_body(
        doc,
        "Aplikasi dikembangkan dengan C++17, dikompilasi menggunakan g++ pada Windows. "
        "Build dilakukan melalui build.bat yang mengompilasi sqlite3.c sebagai objek C, "
        "kemudian menautkannya dengan sumber C++ aplikasi. Program dijalankan sebagai "
        "console application.",
    )

    add_subheading(doc, "4.2 Implementasi Tipe Data dan Variabel")
    add_body(
        doc,
        "Struct Pasien menjadi fondasi seluruh modul. Variabel global pada main.cpp "
        "(arrayPasien, linkedList, antrian, riwayat, treePoli, graph) merepresentasikan "
        "keadaan sistem selama sesi berjalan, sementara database SQLite menjaga persistensi "
        "data pasien antar sesi.",
    )

    add_subheading(doc, "4.3 Implementasi Array dan Linked List")
    add_body(
        doc,
        "ArrayPasien menyediakan akses berindeks, pengecekan kapasitas, serta aksesor "
        "getAt/setAt/tukar untuk mendukung sorting dan searching. LinkedListPasien menyimpan "
        "node dinamis. Keduanya disinkronkan pada operasi tambah, update, dan hapus. Sorting "
        "mengubah urutan Array; Linked List tetap mengikuti urutan insert di head.",
    )

    add_subheading(doc, "4.4 Implementasi Queue dan Stack")
    add_body(
        doc,
        "Menu 7 mengelola antrian: enqueue memasukkan pasien terdaftar ke Queue dan "
        "memperbarui status; dequeue mengeluarkan pasien, mengubah status menjadi Selesai, "
        "mendorong data ke Stack riwayat, serta memperbarui Array dan SQLite. Menu 8 "
        "menampilkan atau mengambil (pop) riwayat terbaru.",
    )

    add_subheading(doc, "4.5 Implementasi Tree")
    add_body(
        doc,
        "Saat program dimulai, treePoli.inisialisasi() membangun hierarki poli. Setiap "
        "penambahan pasien memanggil cariPoli untuk memastikan poli tujuan valid. Menu 9 "
        "menampilkan struktur tree kepada pengguna.",
    )

    add_subheading(doc, "4.6 Implementasi Graph, BFS, dan DFS")
    add_body(
        doc,
        "Graph diinisialisasi dengan edge antar poli: Umum–Gigi, Umum–Anak, Gigi–Penyakit "
        "Dalam, dan Anak–Penyakit Dalam. Menu 13 menampilkan hubungan, menu 14 menjalankan "
        "BFS, dan menu 15 menjalankan DFS. Manfaat e-health: petugas dapat memahami jalur "
        "rujukan antar unit layanan.",
    )

    add_subheading(doc, "4.7 Implementasi Sorting")
    add_body(
        doc,
        "Menu 11 menyediakan Bubble Sort, Selection Sort, dan Insertion Sort dengan pilihan "
        "kriteria No RM, Nama, atau Umur. Setiap algoritma menampilkan jumlah tukar/geser "
        "serta notasi Big O-nya. Manfaatnya: daftar pasien lebih mudah dibaca oleh petugas.",
    )

    add_subheading(doc, "4.8 Implementasi Searching dan Big O")
    add_body(
        doc,
        "Menu 12 membandingkan Linear Search dengan Binary Search berdasarkan jumlah "
        "perbandingan, menampilkan analisis Big O seluruh sistem, serta opsi mengurutkan "
        "No RM terlebih dahulu sebelum binary search. Linear Search O(n); Binary Search "
        "O(log n) dengan syarat data terurut.",
    )

    add_subheading(doc, "4.9 Implementasi Persistensi SQLite")
    add_body(
        doc,
        "Modul Database membuka data/klinik.db, membuat tabel pasien bila belum ada, "
        "menyimpan insert/update/delete, dan memuat seluruh data saat startup ke Array "
        "serta Linked List. Dengan demikian, data tidak hilang dan tidak perlu diinput dua kali.",
    )

    add_subheading(doc, "4.10 Pembahasan Manfaat bagi Pengguna")
    add_bullet(doc, "Petugas dapat mengelola rekam medis dengan validasi yang meminimalkan kesalahan.")
    add_bullet(doc, "Antrian FIFO membuat pelayanan lebih adil dan teratur.")
    add_bullet(doc, "Stack memudahkan peninjauan riwayat pemeriksaan terbaru.")
    add_bullet(doc, "Tree memastikan pasien hanya diarahkan ke poli yang tersedia.")
    add_bullet(doc, "Graph BFS/DFS membantu memahami jaringan rujukan antar poli.")
    add_bullet(doc, "Sorting dan searching mempercepat penyiapan daftar serta pencarian No RM.")
    add_bullet(doc, "SQLite menjaga keberlanjutan data antar sesi kerja.")

    # ========== BAB V ==========
    add_heading_custom(doc, "BAB V", size=14)
    add_heading_custom(doc, "PENUTUP", size=14, space_before=0)

    add_subheading(doc, "5.1 Kesimpulan")
    add_body(
        doc,
        "Sistem Manajemen Antrian dan Rekam Medis Pasien Berbasis E-Health berhasil "
        "mengimplementasikan tipe data dan variabel, Array, Linked List, Queue, Stack, Tree, "
        "Graph, BFS, DFS, sorting, searching, serta analisis Big O pada studi kasus klinik. "
        "Seluruh konsep tersebut terintegrasi dalam menu aplikasi dan bermanfaat bagi pengguna "
        "dalam pengelolaan data pasien, antrian, riwayat, rujukan, serta efisiensi pencarian.",
    )
    add_body(
        doc,
        "Dengan penambahan SQLite, data pasien tersimpan permanen sehingga aplikasi lebih "
        "praktis digunakan berulang kali tanpa kehilangan data.",
    )

    add_subheading(doc, "5.2 Saran")
    add_bullet(doc, "Menyimpan juga state antrian dan riwayat stack ke database.")
    add_bullet(doc, "Mengembangkan antarmuka GUI agar lebih mudah digunakan petugas non-teknis.")
    add_bullet(doc, "Menambahkan algoritma sorting yang lebih efisien (misalnya O(n log n)) untuk data besar.")
    add_bullet(doc, "Memperluas graph rujukan (misalnya Lab dan Apotek) serta fitur jalur terpendek antar unit.")

    # ========== DAFTAR PUSTAKA ==========
    add_heading_custom(doc, "DAFTAR PUSTAKA", size=14)
    refs = [
        "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press.",
        "Goodrich, M. T., Tamassia, R., & Mount, D. M. (2011). Data Structures and Algorithms in C++ (2nd ed.). Wiley.",
        "SQLite Consortium. (2025). SQLite Documentation. https://www.sqlite.org/docs.html",
        "Stroustrup, B. (2013). The C++ Programming Language (4th ed.). Addison-Wesley.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(r)
        set_run_font(run, size=12)


def main():
    doc = Document(SRC)
    keep_cover_only(doc)
    build_content(doc)
    doc.save(OUT)
    print("OK:", OUT)


if __name__ == "__main__":
    main()
